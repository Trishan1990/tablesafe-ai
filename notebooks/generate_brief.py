from docx import Document
from docx.shared import Inches
from datetime import date

# Actual EXP-006 pipeline output
sample_alert = {
    "experiment_id": "EXP-006-LIVE",
    "event": "Moringa Leaf Powder Salmonella Outbreak",
    "product": "Live it Up Super Greens / Moringa-based dietary supplements",
    "pathogen": "Salmonella Typhimurium + Salmonella Newport",
    "recall_date": "2026-01-14",
    "alert_date": "2025-12-15",
    "days_early": 30,
    "result": "PASS",
    "composite_risk_at_alert": 0.7909,
    "peak_composite_risk": 0.9415,
    "f1_score": 1.000,
    "precision": 1.000,
    "recall": 1.000
}

doc = Document()

doc.add_heading("TableSafe AI Evidence Brief", 0)

doc.add_paragraph(f"Experiment ID: {sample_alert['experiment_id']}")
doc.add_paragraph(f"Generated Date: {date.today()}")

doc.add_heading("Event Summary", level=1)
doc.add_paragraph(f"Event: {sample_alert['event']}")
doc.add_paragraph(f"Product: {sample_alert['product']}")
doc.add_paragraph(f"Pathogen: {sample_alert['pathogen']}")

doc.add_heading("Alert Result", level=1)

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Value"

metrics = [
    ("Official Recall Date", sample_alert["recall_date"]),
    ("TableSafe AI Alert Date", sample_alert["alert_date"]),
    ("Early Warning Lead Time", f"{sample_alert['days_early']} days"),
    ("Result", sample_alert["result"]),
    ("Composite Risk at Alert", str(sample_alert["composite_risk_at_alert"])),
    ("Peak Composite Risk", str(sample_alert["peak_composite_risk"])),
    ("F1 Score", str(sample_alert["f1_score"])),
    ("Precision", str(sample_alert["precision"])),
    ("Recall", str(sample_alert["recall"]))
]

for metric, value in metrics:
    row = table.add_row().cells
    row[0].text = metric
    row[1].text = value

doc.add_heading("Signal Sources", level=1)
doc.add_paragraph(
    "The alert was generated from public weak-signal sources including Google Trends, "
    "Reddit discussions, CDC case indicators, FDA outbreak documentation, and OpenFDA regulatory evidence."
)

doc.add_heading("Key Finding", level=1)
doc.add_paragraph(
    "TableSafe AI generated an alert 30 days before the official FDA recall date. "
    "This exceeded the experiment success threshold of 14 days."
)

doc.add_heading("Interpretation", level=1)
doc.add_paragraph(
    "EXP-006 supports the hypothesis that public weak-signal fusion can provide early-warning "
    "indicators for contamination risk before formal regulatory action."
)

doc.add_heading("Recommended Action", level=1)
doc.add_paragraph(
    "Flag moringa-based supplement products for investigation and continue monitoring related "
    "consumer search behavior, illness discussions, and regulatory updates."
)

output_file = "EXP006_TableSafe_AI_Evidence_Brief.docx"
doc.save(output_file)

print(f"Evidence brief generated: {output_file}")